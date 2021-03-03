#!/usr/bin/env python3

"""docx-indent
This script indents docx paragraphs and tables according to a determined
alignment style

Usage:

./docx-indent.py DOCX_PATH
"""

import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Table code: https://github.com/python-openxml/python-docx/issues/209#issuecomment-566128709


def main() -> None:
    # Open the docx document
    document = Document(sys.argv[1])

    # Retrieve all paragraphs present in the document
    paragraphs = document.paragraphs

    # Iterate over each paragraph
    for paragraph in paragraphs:
        # Check whether a paragraph is a title, author name/s, date, and
        # image or table caption
        if paragraph.style.name in [
            "Title",
            "Author",
            "Date",
            "Table Caption",
            "Image Caption",
        ]:
            # Center the caption
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Fully justify the paragraph
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Iterate over each table
    for table in document.tables:
        # Center-align a table
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Allow autofit
        table.allow_autofit = True
        # Autofit the table
        table.autofit = True
        # Set the width of the table to auto, in this way its dimensions
        # will be computed automatically
        table._tblPr.xpath("./w:tblW")[0].attrib[
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
        ] = "auto"
        # Iterate over each row of a table
        for row in table.rows:
            # Iterate over each cell contained in a row
            for cell in row.cells:
                # Get a pointer to the properties of a cell
                tcPr = cell._tc.get_or_add_tcPr()

                # Set the type of a cell according to the content
                tcType = OxmlElement("w:type")
                tcType.set(qn("w:val"), "auto")
                tcPr.append(tcType)

                # Set the width of a cell to 0, since the width will be
                # computed automatically depending on the content
                tcWidth = OxmlElement("w:w")
                tcWidth.set(qn("w:val"), "0")
                tcPr.append(tcWidth)

    # Save the new formatted document
    document.save(sys.argv[1])


if __name__ == "__main__":
    main()
